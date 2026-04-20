"""create_english_template.py — Generate sample_data/english_template.docx.

Run this once after installing dependencies:
    python sample_data/create_english_template.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def add_horizontal_line(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_placeholder(doc, tag, description=""):
    """Add a placeholder paragraph styled as a light grey marker."""
    p = doc.add_paragraph()
    run = p.add_run(tag)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.size = Pt(10)
    run.font.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def main():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Header block ──────────────────────────────────────────────────────────
    title = doc.add_heading("ANNUAL EXAMINATION — 2025", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)

    sub = doc.add_paragraph("Subject: English Grammar   |   Topic: Tenses")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Time Allowed: 2 Hours").bold = True
    info.add_run("          ")
    info.add_run("Maximum Marks: 50").bold = True

    add_horizontal_line(doc)

    inst_head = doc.add_paragraph()
    inst_head.add_run("General Instructions:").bold = True

    instructions = [
        "Read all questions carefully before answering.",
        "Answer all questions. All sections are compulsory.",
        "Write clearly and legibly.",
        "Marks for each question are indicated in brackets.",
    ]
    for inst in instructions:
        p = doc.add_paragraph(inst, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.3)

    add_horizontal_line(doc)
    doc.add_paragraph("")

    # ── SECTION A — MCQ ───────────────────────────────────────────────────────
    sec_a = doc.add_heading("SECTION A — Multiple Choice Questions", level=2)
    sec_a.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note_a = doc.add_paragraph()
    note_a.add_run("Choose the most appropriate answer from the options given.")
    note_a.add_run("  [1 mark each]").bold = True
    note_a.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("Easy level questions:")
    add_placeholder(doc, "{{EASY_MCQ}}")

    doc.add_paragraph("Medium level questions:")
    add_placeholder(doc, "{{MEDIUM_MCQ}}")

    doc.add_paragraph("Hard level questions:")
    add_placeholder(doc, "{{HARD_MCQ}}")

    add_horizontal_line(doc)
    doc.add_paragraph("")

    # ── SECTION B — Short Answer ──────────────────────────────────────────────
    sec_b = doc.add_heading("SECTION B — Short Answer Questions", level=2)
    sec_b.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note_b = doc.add_paragraph()
    note_b.add_run("Answer each question in one or two sentences.")
    note_b.add_run("  [2–3 marks each]").bold = True
    note_b.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("Easy level questions:")
    add_placeholder(doc, "{{EASY_SHORT_ANSWER}}")

    doc.add_paragraph("Medium level questions:")
    add_placeholder(doc, "{{MEDIUM_SHORT_ANSWER}}")

    doc.add_paragraph("Hard level questions:")
    add_placeholder(doc, "{{HARD_SHORT_ANSWER}}")

    add_horizontal_line(doc)
    doc.add_paragraph("")

    # ── SECTION C — Long Answer ───────────────────────────────────────────────
    sec_c = doc.add_heading("SECTION C — Long Answer / Essay Questions", level=2)
    sec_c.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    note_c = doc.add_paragraph()
    note_c.add_run("Answer in detail. Use examples where appropriate.")
    note_c.add_run("  [5–10 marks each]").bold = True
    note_c.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("Easy level questions:")
    add_placeholder(doc, "{{EASY_LONG_ANSWER}}")

    doc.add_paragraph("Medium level questions:")
    add_placeholder(doc, "{{MEDIUM_LONG_ANSWER}}")

    doc.add_paragraph("Hard level questions:")
    add_placeholder(doc, "{{HARD_LONG_ANSWER}}")

    add_horizontal_line(doc)
    doc.add_paragraph("")

    # ── Footer ────────────────────────────────────────────────────────────────
    end = doc.add_paragraph("— End of Question Paper —")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.runs[0].bold = True
    end.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Save ──────────────────────────────────────────────────────────────────
    out = os.path.join(os.path.dirname(__file__), "english_template.docx")
    doc.save(out)
    print(f"English template saved to: {out}")
    print()
    print("Placeholders in this template:")
    placeholders = [
        "{{EASY_MCQ}}", "{{MEDIUM_MCQ}}", "{{HARD_MCQ}}",
        "{{EASY_SHORT_ANSWER}}", "{{MEDIUM_SHORT_ANSWER}}", "{{HARD_SHORT_ANSWER}}",
        "{{EASY_LONG_ANSWER}}", "{{MEDIUM_LONG_ANSWER}}", "{{HARD_LONG_ANSWER}}",
    ]
    for ph in placeholders:
        print(f"  {ph}")
    print()
    print("Use these settings in the app:")
    print("  Topic:    English Grammar")
    print("  Subtopic: Tenses")
    print("  Easy: 3   Medium: 2   Hard: 1")


if __name__ == "__main__":
    main()
