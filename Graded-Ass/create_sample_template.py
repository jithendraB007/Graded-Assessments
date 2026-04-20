"""create_sample_template.py — Generate sample_data/template.docx for testing.

Run once after installing dependencies:
    python create_sample_template.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    doc = Document()

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("ANNUAL EXAMINATION — 2025", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Subject: Mathematics | Topic: Algebra")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Time: 3 Hours                                    Max. Marks: 100")
    doc.add_paragraph(
        "Instructions: Attempt all questions. "
        "Each section carries the marks indicated."
    )
    doc.add_paragraph("")  # blank line

    # ── Section A ─────────────────────────────────────────────────────────────
    doc.add_heading("SECTION A — Multiple Choice Questions", level=2)
    doc.add_paragraph("(Each question carries 1 mark)")
    doc.add_paragraph("")

    doc.add_paragraph("Easy MCQ:")
    doc.add_paragraph("{{EASY_MCQ}}")

    doc.add_paragraph("Medium MCQ:")
    doc.add_paragraph("{{MEDIUM_MCQ}}")

    doc.add_paragraph("Hard MCQ:")
    doc.add_paragraph("{{HARD_MCQ}}")

    # ── Section B ─────────────────────────────────────────────────────────────
    doc.add_heading("SECTION B — Short Answer Questions", level=2)
    doc.add_paragraph("(Each question carries 3 marks)")
    doc.add_paragraph("")

    doc.add_paragraph("Easy Short Answer:")
    doc.add_paragraph("{{EASY_SHORT_ANSWER}}")

    doc.add_paragraph("Medium Short Answer:")
    doc.add_paragraph("{{MEDIUM_SHORT_ANSWER}}")

    doc.add_paragraph("Hard Short Answer:")
    doc.add_paragraph("{{HARD_SHORT_ANSWER}}")

    # ── Section C ─────────────────────────────────────────────────────────────
    doc.add_heading("SECTION C — Long Answer Questions", level=2)
    doc.add_paragraph("(Each question carries 10 marks)")
    doc.add_paragraph("")

    doc.add_paragraph("Easy Long Answer:")
    doc.add_paragraph("{{EASY_LONG_ANSWER}}")

    doc.add_paragraph("Medium Long Answer:")
    doc.add_paragraph("{{MEDIUM_LONG_ANSWER}}")

    doc.add_paragraph("Hard Long Answer:")
    doc.add_paragraph("{{HARD_LONG_ANSWER}}")

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    end = doc.add_paragraph("— End of Question Paper —")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = "sample_data/template.docx"
    doc.save(output_path)
    print(f"Sample template saved to: {output_path}")


if __name__ == "__main__":
    main()
