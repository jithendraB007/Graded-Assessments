"""template_parser.py — Scan a .docx template and extract {{PLACEHOLDER}} records."""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Placeholder parsing constants
# ---------------------------------------------------------------------------

DIFFICULTY_TOKENS = {"EASY", "MEDIUM", "HARD"}

# Maps token(s) found in a placeholder name to internal question_type string.
# Compound tokens (e.g. SHORT_ANSWER) are checked before single tokens.
COMPOUND_TYPE_MAP: dict[str, str] = {
    "SHORT_ANSWER": "short_answer",
    "LONG_ANSWER": "long_answer",
    "SHORTANSWER": "short_answer",
    "LONGANSWER": "long_answer",
}

SINGLE_TYPE_MAP: dict[str, str] = {
    "MCQ": "mcq",
    "SHORT": "short_answer",
    "SA": "short_answer",
    "LONG": "long_answer",
    "LA": "long_answer",
}

PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class PlaceholderRecord:
    """A single recognised {{...}} placeholder found in the template."""
    raw: str               # e.g. "{{EASY_MCQ}}"
    paragraph: Paragraph   # live python-docx Paragraph object
    difficulty: str        # "easy" / "medium" / "hard"
    question_type: str     # "mcq" / "short_answer" / "long_answer"
    count_hint: Optional[int]  # number encoded in placeholder name, or None


@dataclass
class TemplateResult:
    document: Document
    placeholders: list[PlaceholderRecord] = field(default_factory=list)
    unknown_placeholders: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Parsing helpers
# ---------------------------------------------------------------------------

def parse_placeholder(raw: str) -> Optional[dict]:
    """Parse a {{...}} string into difficulty, question_type, and optional count_hint.

    Returns None if the placeholder cannot be mapped to a known difficulty + type.
    """
    # Strip braces and normalise to uppercase
    inner = re.sub(r"[{}]", "", raw).strip().upper()
    parts = inner.split("_")

    difficulty: Optional[str] = None
    q_type: Optional[str] = None
    count_hint: Optional[int] = None

    i = 0
    while i < len(parts):
        token = parts[i]

        # Try two-token compound first (e.g. SHORT_ANSWER)
        if i + 1 < len(parts):
            compound = token + "_" + parts[i + 1]
            if compound in COMPOUND_TYPE_MAP and q_type is None:
                q_type = COMPOUND_TYPE_MAP[compound]
                i += 2
                continue

        if token in DIFFICULTY_TOKENS and difficulty is None:
            difficulty = token.lower()
        elif token in SINGLE_TYPE_MAP and q_type is None:
            q_type = SINGLE_TYPE_MAP[token]
        elif token.isdigit() and count_hint is None:
            count_hint = int(token)

        i += 1

    if difficulty is None or q_type is None:
        return None

    return {
        "raw": raw,
        "difficulty": difficulty,
        "question_type": q_type,
        "count_hint": count_hint,
    }


def normalise_placeholder_paragraph(para: Paragraph) -> None:
    """Collapse all runs in a placeholder paragraph into the first run.

    Word often splits a {{...}} tag across multiple runs due to autocorrect,
    spell-check, or partial formatting. This function merges them so the
    placeholder appears as one continuous text string in a single run.
    Only paragraphs that contain a {{...}} pattern are modified.
    """
    full_text = "".join(r.text for r in para.runs)
    if not PLACEHOLDER_RE.search(full_text):
        return
    if not para.runs:
        return
    para.runs[0].text = full_text
    for run in para.runs[1:]:
        run.text = ""


def _collect_paragraphs(doc: Document) -> list[Paragraph]:
    """Return all paragraphs in the document: top-level + inside table cells."""
    paras: list[Paragraph] = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    return paras


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_template(file_buffer: io.BytesIO) -> TemplateResult:
    """Open a .docx template and extract all recognised {{PLACEHOLDER}} records.

    Args:
        file_buffer: BytesIO-compatible buffer of the .docx file.

    Returns:
        TemplateResult containing the live Document, recognised PlaceholderRecords,
        any unrecognised placeholder strings, and warning messages.
    """
    doc = Document(file_buffer)
    result = TemplateResult(document=doc)

    all_paragraphs = _collect_paragraphs(doc)

    for idx, para in enumerate(all_paragraphs):
        # Build full text from runs (more reliable than para.text for split runs)
        full_text = "".join(r.text for r in para.runs) or para.text
        matches = PLACEHOLDER_RE.findall(full_text)
        if not matches:
            continue

        # Merge split runs so replacement works cleanly
        normalise_placeholder_paragraph(para)

        for match in matches:
            raw = "{{" + match + "}}"
            parsed = parse_placeholder(raw)

            if parsed is None:
                result.unknown_placeholders.append(raw)
                result.warnings.append(
                    f"Unrecognised placeholder '{raw}' at paragraph {idx} — "
                    f"it will be left untouched in the document."
                )
            else:
                result.placeholders.append(
                    PlaceholderRecord(
                        raw=raw,
                        paragraph=para,
                        difficulty=parsed["difficulty"],
                        question_type=parsed["question_type"],
                        count_hint=parsed["count_hint"],
                    )
                )

    if not result.placeholders and not result.unknown_placeholders:
        result.warnings.append(
            "No {{PLACEHOLDER}} markers found in the template. "
            "The document will be returned unchanged. "
            "Check the sidebar for the correct placeholder naming convention."
        )

    return result
