"""paper_generator.py — Replace placeholder paragraphs with exam-format question tables.

Table structure (5 columns, matching the standard exam paper format):
  Col 0  — Q.No   (narrow, shaded header column)
  Col 1  — Question content  (wide)
  Col 2  — Mark   (narrow, centred)
  Col 3  — BTL    (narrow, centred — Bloom's Taxonomy Level)
  Col 4  — CO     (narrow, centred — Course Outcome)

Global question numbering runs across ALL placeholders in document order.
Long-answer questions are paired as (a) / (b) with an (OR) merged row between them.

Tables are inserted at XML level (addprevious) so they sit exactly where the
placeholder was, then the placeholder paragraph is removed.
No run properties are ever copied from the placeholder — it is intentionally
styled light-grey/italic and that must not bleed into question text.

When a CourseConfig is provided:
  - build_branding_header() inserts a 2-row header table at the top of the document.
  - build_section_header_rows() prepends PART A/B/C section header rows to each table.
"""
from __future__ import annotations

import io
import os
from typing import Optional

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

from modules.template_parser import TemplateResult, PlaceholderRecord
from modules.question_selector import SelectionResult
from modules.course_config import CourseConfig, SectionConfig, BrandingConfig


# ---------------------------------------------------------------------------
# Layout constants
# ---------------------------------------------------------------------------

# Q.No | Question | Mark | BTL | CO
_COL_WIDTHS = (Inches(0.50), Inches(4.50), Inches(0.65), Inches(0.65), Inches(0.65))

_HEADER_BG   = "1F497D"   # dark navy  — table column-header row
_QNUM_BG     = "D9E1F2"   # light blue — Q.No cell
_OR_BG       = "F2F2F2"   # light grey — (OR) row

_FONT_NAME   = "Calibri"
_FONT_BODY   = Pt(11)
_FONT_SMALL  = Pt(10)
_FONT_OR     = Pt(10)

_ANSWER_LINE = "_" * 60


# ---------------------------------------------------------------------------
# Low-level XML / cell helpers
# ---------------------------------------------------------------------------

def _set_cell_shade(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_col_width(table: Table, col_index: int, width) -> None:
    for row in table.rows:
        cell = row.cells[col_index]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        twips = int(width / 914400 * 1440)
        tcW.set(qn("w:w"), str(twips))
        tcW.set(qn("w:type"), "dxa")


def _clear_cell(cell) -> None:
    tc = cell._tc
    for p in tc.findall(qn("w:p"))[1:]:
        tc.remove(p)
    first_p = tc.find(qn("w:p"))
    if first_p is not None:
        for r in first_p.findall(qn("w:r")):
            first_p.remove(r)


def _write_cell(cell, text: str, bold: bool = False, italic: bool = False,
                color_rgb: tuple = (0x1A, 0x1A, 0x1A),
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                size: Pt = None) -> None:
    """Write text into the first paragraph of a cell with clean formatting."""
    _clear_cell(cell)
    para = cell.paragraphs[0]
    para.alignment = alignment
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = _FONT_NAME
    run.font.size = size or _FONT_BODY
    run.font.color.rgb = RGBColor(*color_rgb)


def _append_cell_para(cell, text: str, bold: bool = False, italic: bool = False,
                      color_rgb: tuple = (0x1A, 0x1A, 0x1A),
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      size: Pt = None) -> None:
    """Append a new paragraph to a cell."""
    para = cell.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = _FONT_NAME
    run.font.size = size or _FONT_BODY
    run.font.color.rgb = RGBColor(*color_rgb)


def _safe_str(row: pd.Series, col: str, fallback: str = "") -> str:
    val = row.get(col, fallback)
    if pd.isna(val) or str(val).strip() in ("", "nan"):
        return fallback
    return str(val).strip()


def _safe_option(row: pd.Series, col: str) -> str:
    return _safe_str(row, col, "[option not provided]")


def _marks_str(row: pd.Series) -> str:
    val = row.get("marks", "")
    if pd.isna(val) or str(val).strip() in ("", "nan"):
        return ""
    try:
        return str(int(float(str(val))))
    except (ValueError, TypeError):
        return str(val).strip()


# ---------------------------------------------------------------------------
# Column-header row
# ---------------------------------------------------------------------------

def _add_header_row(table: Table) -> None:
    """Add the top column-header row: Q.No | Question | Mark | BTL | CO."""
    hdr = table.add_row()
    labels = ["Q.No", "Question", "Mark", "BTL", "CO"]
    aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ]
    for cell, label, align in zip(hdr.cells, labels, aligns):
        _write_cell(cell, label, bold=True, color_rgb=(0xFF, 0xFF, 0xFF),
                    alignment=align, size=_FONT_BODY)
        _set_cell_shade(cell, _HEADER_BG)


# ---------------------------------------------------------------------------
# (OR) separator row
# ---------------------------------------------------------------------------

def _add_or_row(table: Table) -> None:
    """Add a full-width (OR) row spanning all 5 columns."""
    tr = table.add_row()
    # Merge all 5 cells
    merged = tr.cells[0].merge(tr.cells[4])
    _write_cell(merged, "(OR)", bold=True, italic=False,
                color_rgb=(0x44, 0x44, 0x44),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                size=_FONT_OR)
    _set_cell_shade(merged, _OR_BG)


# ---------------------------------------------------------------------------
# Per-question-type row builders
# ---------------------------------------------------------------------------

def _add_mcq_row(table: Table, label: str, row: pd.Series) -> None:
    tr = table.add_row()

    # Q.No
    _write_cell(tr.cells[0], str(label), bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shade(tr.cells[0], _QNUM_BG)

    # Question: bold stem + 2-per-row options
    c = tr.cells[1]
    _clear_cell(c)
    stem_para = c.paragraphs[0]
    stem_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    stem_run = stem_para.add_run(str(row["question_text"]))
    stem_run.bold = True
    stem_run.font.name = _FONT_NAME
    stem_run.font.size = _FONT_BODY
    stem_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    labels_opt = ["a", "b", "c", "d"]
    cols_opt   = ["option_a", "option_b", "option_c", "option_d"]
    for i in range(0, 4, 2):
        left  = f"({labels_opt[i]}) {_safe_option(row, cols_opt[i])}"
        right = f"({labels_opt[i+1]}) {_safe_option(row, cols_opt[i+1])}"
        _append_cell_para(c, f"  {left:<42}  {right}", size=_FONT_SMALL)

    # Mark
    _write_cell(tr.cells[2], _marks_str(row),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # BTL
    _write_cell(tr.cells[3], _safe_str(row, "btl"),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # CO
    _write_cell(tr.cells[4], _safe_str(row, "co"),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _add_short_answer_row(table: Table, label: str, row: pd.Series) -> None:
    tr = table.add_row()

    _write_cell(tr.cells[0], str(label), bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shade(tr.cells[0], _QNUM_BG)

    c = tr.cells[1]
    _clear_cell(c)
    q_para = c.paragraphs[0]
    q_run = q_para.add_run(str(row["question_text"]))
    q_run.bold = True
    q_run.font.name = _FONT_NAME
    q_run.font.size = _FONT_BODY
    q_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    for _ in range(3):
        _append_cell_para(c, _ANSWER_LINE, size=_FONT_SMALL)

    _write_cell(tr.cells[2], _marks_str(row),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tr.cells[3], _safe_str(row, "btl"),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tr.cells[4], _safe_str(row, "co"),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _add_long_answer_row(table: Table, label: str, row: pd.Series) -> None:
    tr = table.add_row()

    _write_cell(tr.cells[0], str(label), bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_shade(tr.cells[0], _QNUM_BG)

    c = tr.cells[1]
    _clear_cell(c)
    q_para = c.paragraphs[0]
    q_run = q_para.add_run(str(row["question_text"]))
    q_run.bold = True
    q_run.font.name = _FONT_NAME
    q_run.font.size = _FONT_BODY
    q_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    for _ in range(7):
        _append_cell_para(c, _ANSWER_LINE, size=_FONT_SMALL)

    _write_cell(tr.cells[2], _marks_str(row),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tr.cells[3], _safe_str(row, "btl"),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell(tr.cells[4], _safe_str(row, "co"),
                alignment=WD_ALIGN_PARAGRAPH.CENTER)


def _add_note_row(table: Table, note: str) -> None:
    """Full-width italic note row (e.g. shortfall warning)."""
    tr = table.add_row()
    merged = tr.cells[0].merge(tr.cells[4])
    _write_cell(merged, note, italic=True,
                color_rgb=(0x66, 0x66, 0x66), size=_FONT_SMALL)


# ---------------------------------------------------------------------------
# Table factory
# ---------------------------------------------------------------------------

def _build_question_table(
    doc: Document,
    df: pd.DataFrame,
    question_type: str,
    start_q: int,           # global question counter value BEFORE this table
    note: Optional[str],
) -> tuple[Table, int]:
    """Create a question table and return it with the updated global counter.

    For MCQ / short_answer: each row increments the counter by 1.
    For long_answer: questions come in (a)/(b) pairs; counter increments by 1
    per *pair* (not per question), matching the exam format 21(a)/21(b).

    Returns:
        (table, new_counter)  where new_counter is the counter value after
        all questions in this table have been numbered.
    """
    table = doc.add_table(rows=0, cols=5)
    table.style = "Table Grid"

    _add_header_row(table)

    qt = question_type.lower()
    counter = start_q
    pair_base = counter  # reused inside the long-answer loop

    if qt == "long_answer":
        for idx, (_, row) in enumerate(df.iterrows()):
            if idx % 2 == 0:                    # first of a new pair → (a)
                counter += 1
                pair_base = counter
                _add_long_answer_row(table, f"{pair_base}(a)", row)
            else:                               # second of pair → OR + (b)
                _add_or_row(table)
                _add_long_answer_row(table, f"{pair_base}(b)", row)
    else:
        for _, row in df.iterrows():
            counter += 1
            if qt == "mcq":
                _add_mcq_row(table, counter, row)
            elif qt == "short_answer":
                _add_short_answer_row(table, counter, row)
            else:
                # Generic fallback
                tr = table.add_row()
                _write_cell(tr.cells[0], str(counter), bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_shade(tr.cells[0], _QNUM_BG)
                _write_cell(tr.cells[1], str(row["question_text"]))
                _write_cell(tr.cells[2], _marks_str(row),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
                _write_cell(tr.cells[3], _safe_str(row, "btl"),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
                _write_cell(tr.cells[4], _safe_str(row, "co"),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if note:
        _add_note_row(table, note)

    # Apply column widths
    for col_idx, width in enumerate(_COL_WIDTHS):
        _set_col_width(table, col_idx, width)

    return table, counter


# ---------------------------------------------------------------------------
# XML-level insertion
# ---------------------------------------------------------------------------

def _insert_table_before_paragraph(table: Table, para) -> None:
    """Relocate table's <w:tbl> to just before the placeholder paragraph,
    then remove the placeholder paragraph."""
    tbl_elem = table._tbl
    p_elem   = para._p
    parent   = p_elem.getparent()
    if parent is None:
        return
    tbl_parent = tbl_elem.getparent()
    if tbl_parent is not None:
        tbl_parent.remove(tbl_elem)
    p_elem.addprevious(tbl_elem)
    parent.remove(p_elem)


def _replace_para_with_note(para, note: str) -> None:
    """Swap the placeholder paragraph for a plain italic note paragraph."""
    p_elem = para._p
    parent = p_elem.getparent()
    if parent is None:
        return
    new_p = OxmlElement("w:p")
    r     = OxmlElement("w:r")
    rPr   = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:i"))
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = note
    r.append(t)
    new_p.append(r)
    p_elem.addprevious(new_p)
    parent.remove(p_elem)


# ---------------------------------------------------------------------------
# Branding header (inserted before all section tables)
# ---------------------------------------------------------------------------

_SECTION_HDR_BG  = "2E4057"   # darker navy for PART A/B/C header rows
_SECTION_SUB_BG  = "D6E4F0"   # pale blue for sub-header (question type label)


def build_branding_header(
    doc: Document,
    branding: BrandingConfig,
    exam_title: str = "ANNUAL EXAMINATION",
    time_allowed: str = "3 Hours",
    total_marks: float = 100,
) -> None:
    """Insert a branded 2-row header table at the very start of the document body.

    Row 1: [Logo cell] | [Institution name / dept / regulation (merged with remaining cols)]
    Row 2 (merged): [Exam title  |  Time: X  |  Max Marks: Y]
    """
    body = doc.element.body
    first_child = body[0] if len(body) else None

    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"

    # --- Row 0: logo | institution details ---
    r0 = tbl.rows[0]
    logo_cell = r0.cells[0]
    info_cell = r0.cells[1].merge(r0.cells[2])

    # Logo
    if branding.logo_path and os.path.exists(branding.logo_path):
        _clear_cell(logo_cell)
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = logo_para.add_run()
            run.add_picture(branding.logo_path, width=Inches(0.9))
        except Exception:
            _write_cell(logo_cell, "[Logo]", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        size=Pt(9))
    else:
        _write_cell(logo_cell, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Institution block
    _clear_cell(info_cell)
    ip = info_cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_run = ip.add_run(branding.institution_name)
    inst_run.bold = True
    inst_run.font.size = Pt(14)
    inst_run.font.name = _FONT_NAME
    inst_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    if branding.department:
        _append_cell_para(info_cell, branding.department, bold=False,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11))
    if branding.regulation:
        _append_cell_para(info_cell, branding.regulation, bold=False, italic=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10),
                          color_rgb=(0x44, 0x44, 0x44))

    # Column widths: logo narrow, info wide
    for row in tbl.rows:
        for ci, width in enumerate([Inches(1.0), Inches(3.0), Inches(2.45)]):
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(width / 914400 * 1440)))
            tcW.set(qn("w:type"), "dxa")
            existing = tcPr.find(qn("w:tcW"))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcW)

    # --- Row 1: exam title / time / marks (merged) ---
    r1 = tbl.rows[1]
    merged = r1.cells[0].merge(r1.cells[2])
    banner = f"{exam_title}    |    Time Allowed: {time_allowed}    |    Maximum Marks: {int(total_marks)}"
    _write_cell(merged, banner, bold=True,
                color_rgb=(0xFF, 0xFF, 0xFF),
                alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11))
    _set_cell_shade(merged, _HEADER_BG)

    # Move the table to the top of body (before everything else)
    tbl_elem = tbl._tbl
    tbl_parent = tbl_elem.getparent()
    if tbl_parent is not None:
        tbl_parent.remove(tbl_elem)
    if first_child is not None:
        first_child.addprevious(tbl_elem)
    else:
        body.append(tbl_elem)

    # Spacer paragraph after header
    spacer = OxmlElement("w:p")
    tbl_elem.addnext(spacer)


# ---------------------------------------------------------------------------
# Section header rows (PART A / B / C banner inside each section table)
# ---------------------------------------------------------------------------

def build_section_header_rows(
    table: Table,
    section: SectionConfig,
    actual_q_count: int,
) -> None:
    """Prepend two merged header rows at the top of a section table.

    Row 0 (dark navy, white text):
        "PART A  (20×1 = 20 Marks)  Answer all the Questions"
    Row 1 (light blue):
        "(Multiple Choice Questions)"

    *actual_q_count* is used instead of the configured count so that after
    judge-based rebalancing the marks total is still accurate.
    """
    marks_per_q = section.marks_per_q
    if section.either_or:
        pairs = actual_q_count // 2
        total = pairs * marks_per_q
        count_label = f"{pairs}×{int(marks_per_q)} = {int(total)}"
    else:
        total = actual_q_count * marks_per_q
        count_label = f"{actual_q_count}×{int(marks_per_q)} = {int(total)}"

    part_banner = (
        f"{section.label}  ({count_label} Marks)  {section.instruction}"
    )

    # We must insert rows at position 0 (before the column header row).
    # python-docx Table has no insert_row_at(), so we manipulate XML directly.
    tbl_elem = table._tbl

    def _make_merged_row(text: str, bg_hex: str, text_rgb: tuple, bold: bool) -> None:
        tr = OxmlElement("w:tr")
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        # gridSpan = 5
        gridSpan = OxmlElement("w:gridSpan")
        gridSpan.set(qn("w:val"), "5")
        tcPr.append(gridSpan)
        # shading
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        tcPr.append(shd)
        tc.append(tcPr)
        # paragraph
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "%02X%02X%02X" % text_rgb)
        rPr.append(color_el)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")   # 11pt
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
        return tr

    # Build the two rows
    row_banner = _make_merged_row(
        part_banner, _SECTION_HDR_BG, (0xFF, 0xFF, 0xFF), bold=True
    )
    row_sub = _make_merged_row(
        f"({section.title})", _SECTION_SUB_BG, (0x1A, 0x1A, 0x1A), bold=False
    )

    # Insert before the first existing <w:tr> (the column-header row)
    first_tr = tbl_elem.find(qn("w:tr"))
    if first_tr is not None:
        first_tr.addprevious(row_sub)
        row_sub.addprevious(row_banner)
    else:
        tbl_elem.append(row_banner)
        tbl_elem.append(row_sub)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_paper(
    template_result: TemplateResult,
    selection_result: SelectionResult,
    course_config: Optional[CourseConfig] = None,
) -> io.BytesIO:
    """Replace all placeholder paragraphs with formatted 5-column question tables.

    Question numbers are global — they run continuously across all placeholders
    in document order.  Long-answer questions are paired as N(a) / (OR) / N(b).

    When *course_config* is provided:
    - A branded institution header is prepended to the document.
    - Each section table gets PART A/B/C merged header rows showing marks totals.
    - Section header rows are matched to SectionConfig by question_type.

    Args:
        template_result:   From template_parser.parse_template().
        selection_result:  From question_selector.select_questions().
        course_config:     Optional; enables branding + section headers.

    Returns:
        BytesIO buffer of the ready-to-download .docx file.
    """
    doc = template_result.document
    global_counter = 0   # increments as questions are numbered

    # Pre-compute total question counts per section (for accurate marks totals in headers).
    # Key: SectionConfig.part letter  →  total assigned question count for that section.
    _section_total: dict[str, int] = {}
    _section_header_done: set[str] = set()   # parts whose header row has been emitted

    if course_config:
        for record in template_result.placeholders:
            df = selection_result.assignments.get(record.raw, pd.DataFrame())
            qt = record.question_type.lower()
            sec = next(
                (s for s in course_config.sections if s.question_type.lower() == qt),
                None,
            )
            if sec:
                _section_total[sec.part] = _section_total.get(sec.part, 0) + len(df)

    for record in template_result.placeholders:
        assigned_df = selection_result.assignments.get(record.raw, pd.DataFrame())

        error = next(
            (e for e in selection_result.errors if e.placeholder_raw == record.raw),
            None,
        )

        if assigned_df.empty:
            note = (
                f"[NOTE: 0 of {error.requested} requested questions available "
                f"for {record.raw} — check topic / subtopic / difficulty in your dataset.]"
                if error else ""
            )
            _replace_para_with_note(record.paragraph, note)
            continue

        shortfall_note = (
            f"[NOTE: Only {error.available} of {error.requested} questions available "
            f"for {record.raw}.]"
            if error else None
        )

        table, global_counter = _build_question_table(
            doc,
            assigned_df,
            record.question_type,
            global_counter,
            shortfall_note,
        )

        # --- Optionally prepend section header rows (first table of each part only) ---
        if course_config:
            qt = record.question_type.lower()
            sec = next(
                (s for s in course_config.sections if s.question_type.lower() == qt),
                None,
            )
            if sec and sec.part not in _section_header_done:
                total_q = _section_total.get(sec.part, len(assigned_df))
                build_section_header_rows(table, sec, total_q)
                _section_header_done.add(sec.part)

        _insert_table_before_paragraph(table, record.paragraph)

    # --- Optionally prepend branded institution header ---
    if course_config:
        build_branding_header(
            doc,
            course_config.branding,
            exam_title=course_config.exam_title,
            time_allowed=course_config.time_allowed,
            total_marks=course_config.grand_total_marks,
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
