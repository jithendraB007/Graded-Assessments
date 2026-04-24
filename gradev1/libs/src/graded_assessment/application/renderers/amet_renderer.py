from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from graded_assessment.application.renderers._base import insert_logo, open_template, set_table_borders
from graded_assessment.domain.amet_types import AmetAssessmentRequest


def _bold_center(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_table_row(table, cells: list[str], bold: bool = False) -> None:
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text
        if bold:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


def _merge_row(table, row_idx: int, text: str, bold: bool = True) -> None:
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[-1])
    row.cells[0].text = text
    if bold:
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def render(request: AmetAssessmentRequest) -> bytes:
    doc = open_template("AMET")
    insert_logo(doc, "amet")

    # ── Header ──────────────────────────────────────────────────
    _bold_center(doc, request.exam_type, size=13)
    doc.add_paragraph(
        f"Programme & Batch: {request.programme}"
        f"{'':>20}Semester : {request.semester}"
    )
    doc.add_paragraph(
        f"Course Name : {request.course_name}"
        f"{'':>20}Course Code : {request.course_code}"
    )
    doc.add_paragraph(
        f"Duration : {request.duration}"
        f"{'':>30}Maximum Marks: {request.max_marks} marks"
    )
    doc.add_paragraph()

    # ── Instructions ────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run("Instructions:").bold = True
    for i, instr in enumerate(request.instructions, start=1):
        doc.add_paragraph(f"{i}. {instr}")

    doc.add_paragraph()

    # ── Question table ──────────────────────────────────────────
    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Question No", "Question", "Mark", "BTL", "CO"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    # Part A
    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = (
        f"PART A ({request.part_a.total}) "
        f"{request.part_a.instruction}"
    )
    row.cells[0].paragraphs[0].runs[0].bold = True

    for q in request.part_a.questions:
        _add_table_row(table, [q.number, q.text, str(q.mark), q.btl, q.co])

    # Part B
    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = (
        f"PART B ({request.part_b.total}) "
        f"{request.part_b.instruction}"
    )
    row.cells[0].paragraphs[0].runs[0].bold = True

    for pair in request.part_b.question_pairs:
        _add_table_row(table, [pair.a.number, pair.a.text, str(pair.a.mark), pair.a.btl, pair.a.co])
        row = table.add_row()
        row.cells[0].merge(row.cells[4])
        row.cells[0].text = "(OR)"
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_table_row(table, [pair.b.number, pair.b.text, str(pair.b.mark), pair.b.btl, pair.b.co])

    # Part C
    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = (
        f"PART C ({request.part_c.total}) "
        f"{request.part_c.instruction}"
    )
    row.cells[0].paragraphs[0].runs[0].bold = True

    q = request.part_c.question
    _add_table_row(table, [q.number, q.text, str(q.mark), q.btl, q.co])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
