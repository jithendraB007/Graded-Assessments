from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from graded_assessment.domain.svyasa_types import SvyasaAssessmentRequest


def render(request: SvyasaAssessmentRequest) -> bytes:
    doc = Document()

    # ── USN row ──────────────────────────────────────────────────
    usn_table = doc.add_table(rows=1, cols=12)
    usn_table.style = "Table Grid"
    usn_table.rows[0].cells[0].text = "USN"
    doc.add_paragraph()

    # ── Header table ─────────────────────────────────────────────
    hdr_table = doc.add_table(rows=5, cols=4)
    hdr_table.style = "Table Grid"
    hdr_data = [
        ["Month & Year of Examination", request.month_year, "Academic year", request.academic_year],
        ["Program", request.program, "Specialization", request.specialization],
        ["Semester", request.semester, "Date of Examination", request.date_of_exam],
        ["Course Code", request.course_code, request.course_code, request.course_code],
        ["Course Name", request.course_name, request.course_name, request.course_name],
    ]
    for r_idx, row_data in enumerate(hdr_data):
        for c_idx, text in enumerate(row_data):
            hdr_table.rows[r_idx].cells[c_idx].text = text

    doc.add_paragraph()

    # ── Part A ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Part – A")
    run.bold = True

    part_a_count = len(request.part_a.questions)
    marks_each = request.part_a.questions[0].marks if request.part_a.questions else 3
    doc.add_paragraph(
        f"Answer all the questions   "
        f"{part_a_count} Q x {marks_each} M = {part_a_count * marks_each}"
    )

    table_a = doc.add_table(rows=1, cols=5)
    table_a.style = "Table Grid"
    for i, h in enumerate(["Q.No.", "Questions", "CO", "RBTL", "Marks"]):
        table_a.rows[0].cells[i].text = h
        table_a.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for q in request.part_a.questions:
        row = table_a.add_row()
        row.cells[0].text = q.number
        row.cells[1].text = q.text
        row.cells[2].text = q.co
        row.cells[3].text = q.rbtl
        row.cells[4].text = str(q.marks)

    doc.add_paragraph()

    # ── Part B ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Part - B")
    run.bold = True

    pair_count = len(request.part_b.question_pairs)
    marks_each_b = request.part_b.question_pairs[0].a.marks if request.part_b.question_pairs else 14
    doc.add_paragraph(
        f"Answer all the questions   "
        f"{pair_count} Q x {marks_each_b} M = {pair_count * marks_each_b}"
    )

    table_b = doc.add_table(rows=1, cols=5)
    table_b.style = "Table Grid"
    for i, h in enumerate(["Q.No.", "Questions", "CO", "RBTL", "Marks"]):
        table_b.rows[0].cells[i].text = h
        table_b.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for pair in request.part_b.question_pairs:
        row_a = table_b.add_row()
        row_a.cells[0].text = pair.a.number
        row_a.cells[1].text = pair.a.text
        row_a.cells[2].text = pair.a.co
        row_a.cells[3].text = pair.a.rbtl
        row_a.cells[4].text = str(pair.a.marks)

        or_row = table_b.add_row()
        or_row.cells[0].merge(or_row.cells[4])
        or_row.cells[0].text = "OR"
        or_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_b = table_b.add_row()
        row_b.cells[0].text = pair.b.number
        row_b.cells[1].text = pair.b.text
        row_b.cells[2].text = pair.b.co
        row_b.cells[3].text = pair.b.rbtl
        row_b.cells[4].text = str(pair.b.marks)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
