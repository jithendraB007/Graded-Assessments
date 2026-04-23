from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from graded_assessment.domain.cdu_types import CduAssessmentRequest, CduSet


def _render_set(doc: Document, cdu_set: CduSet, university_name: str, course_info: str, time: str, max_marks: int) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    # Set header row
    hdr_row = table.add_row()
    hdr_row.cells[0].merge(hdr_row.cells[1])
    hdr_row.cells[0].text = (
        f"{cdu_set.label}\n{university_name}\n{course_info}"
    )
    hdr_row.cells[0].paragraphs[0].runs[0].bold = True
    hdr_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    time_row = table.add_row()
    time_row.cells[0].text = f"Time: {time}]"
    time_row.cells[1].text = f"[Max. Marks: {max_marks}"

    # Section A
    sec_a_row = table.add_row()
    sec_a_row.cells[0].merge(sec_a_row.cells[1])
    sec_a_row.cells[0].text = "Section - A"
    sec_a_row.cells[0].paragraphs[0].runs[0].bold = True

    instr_row = table.add_row()
    instr_row.cells[0].merge(instr_row.cells[1])
    instr_row.cells[0].text = cdu_set.section_a.instruction

    for q in cdu_set.section_a.questions:
        row = table.add_row()
        row.cells[0].text = q.number
        row.cells[1].text = q.text

    # Section B
    sec_b_row = table.add_row()
    sec_b_row.cells[0].merge(sec_b_row.cells[1])
    sec_b_row.cells[0].text = "Section - B"
    sec_b_row.cells[0].paragraphs[0].runs[0].bold = True

    instr_b_row = table.add_row()
    instr_b_row.cells[0].merge(instr_b_row.cells[1])
    instr_b_row.cells[0].text = cdu_set.section_b.instruction

    for pair in cdu_set.section_b.question_pairs:
        row_a = table.add_row()
        row_a.cells[0].text = pair.a.number
        row_a.cells[1].text = pair.a.text

        or_row = table.add_row()
        or_row.cells[0].merge(or_row.cells[1])
        or_row.cells[0].text = "OR"
        or_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_b = table.add_row()
        row_b.cells[0].text = pair.b.number
        row_b.cells[1].text = pair.b.text


def render(request: CduAssessmentRequest) -> bytes:
    doc = Document()

    for i, cdu_set in enumerate(request.sets):
        _render_set(
            doc, cdu_set,
            request.university_name, request.course_info,
            request.time, request.max_marks,
        )
        if i < len(request.sets) - 1:
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
