from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from graded_assessment.application.renderers._base import insert_logo, open_template, set_table_borders
from graded_assessment.domain.anu_types import AnuAssessmentRequest


def _bold_center(doc: Document, text: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def render(request: AnuAssessmentRequest) -> bytes:
    doc = open_template("ANU")
    insert_logo(doc, "anu")

    # ── Header ──────────────────────────────────────────────────
    _bold_center(doc, request.university_name, size=14)
    _bold_center(doc, request.batch)
    _bold_center(doc, request.exam_type)
    _bold_center(doc, request.course_name)
    doc.add_paragraph(
        f"Date: {request.date}   "
        f"Duration: {request.duration}   "
        f"Max. Marks: {request.max_marks}"
    )
    doc.add_paragraph()

    # ── Hall Ticket row ─────────────────────────────────────────
    ht_table = doc.add_table(rows=1, cols=12)
    set_table_borders(ht_table)
    ht_table.rows[0].cells[0].text = "H.T. No:-"

    doc.add_paragraph()

    # ── Notes ───────────────────────────────────────────────────
    for i, note in enumerate(request.notes, start=1):
        doc.add_paragraph(f"{i}. {note}")
    doc.add_paragraph()

    # ── Part A table ────────────────────────────────────────────
    part_a_table = doc.add_table(rows=1, cols=6)
    set_table_borders(part_a_table)

    hdr = part_a_table.rows[0]
    for i, h in enumerate(["PART-A", "", "", "", "Course Outcomes", "Bloom's level"]):
        hdr.cells[i].text = h
        if h:
            hdr.cells[i].paragraphs[0].runs[0].bold = True

    hdr.cells[0].merge(hdr.cells[3])

    instr_row = part_a_table.add_row()
    instr_row.cells[0].merge(instr_row.cells[3])
    instr_row.cells[0].text = "Answer all the following short answer questions."

    q_row = part_a_table.add_row()
    q_row.cells[0].text = "1"
    for sq in request.part_a.sub_questions:
        sub_row = part_a_table.add_row()
        sub_row.cells[1].text = sq.sub
        sub_row.cells[2].merge(sub_row.cells[3])
        sub_row.cells[2].text = sq.text
        sub_row.cells[4].text = sq.co
        sub_row.cells[5].text = sq.bloom

    doc.add_paragraph()

    # ── Part B table ────────────────────────────────────────────
    part_b_table = doc.add_table(rows=1, cols=6)
    set_table_borders(part_b_table)

    hdr_b = part_b_table.rows[0]
    for i, h in enumerate(["PART B", "", "", "Marks", "Course Outcomes", "Bloom's Level"]):
        hdr_b.cells[i].text = h
        if h:
            hdr_b.cells[i].paragraphs[0].runs[0].bold = True

    for q in request.part_b.questions:
        row = part_b_table.add_row()
        row.cells[0].text = q.number
        row.cells[2].merge(row.cells[1])
        row.cells[1].text = q.text
        row.cells[3].text = q.marks
        row.cells[4].text = q.co
        row.cells[5].text = q.bloom

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
