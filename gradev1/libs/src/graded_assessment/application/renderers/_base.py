from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

ASSETS_DIR = Path(__file__).parents[5] / "assets"
TEMPLATES_DIR = ASSETS_DIR / "templates"
LOGOS_DIR = ASSETS_DIR / "logos"


def open_template(university_id: str) -> Document:
    """Open the university .docx template and clear its content, keeping styles."""
    template_path = TEMPLATES_DIR / f"{university_id}.docx"
    if template_path.exists():
        doc = Document(str(template_path))
        # Clear all body content while preserving section properties (margins, page size)
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "sectPr":
                body.remove(child)
    else:
        doc = Document()
    return doc


def set_table_borders(table) -> None:
    """Apply visible borders to every cell in a table (works regardless of template styles)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def insert_logo(doc: Document, university_id: str, width_inches: float = 1.8) -> None:
    """Insert university logo centred at the current position if the file exists."""
    logo_path = LOGOS_DIR / f"{university_id}.png"
    if not logo_path.exists():
        # Try common image extensions
        for ext in (".jpg", ".jpeg", ".PNG", ".JPG"):
            alt = LOGOS_DIR / f"{university_id}{ext}"
            if alt.exists():
                logo_path = alt
                break
        else:
            return  # No logo found — skip silently

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(logo_path), width=Inches(width_inches))
