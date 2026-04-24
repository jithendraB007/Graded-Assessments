"""
AMET Graded Assessment Generator
---------------------------------
Self-contained script. All assets are resolved relative to this file.

  assets/templates/AMET.docx  — university page layout template
  assets/logos/amet.png       — university logo

Usage:
    from generate import generate, AmetAssessmentRequest, ...
    data = AmetAssessmentRequest(...)
    path = generate(data)          # returns saved file path
"""

from __future__ import annotations

import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pydantic import BaseModel

# ── Paths ────────────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent
TEMPLATE_PATH = _HERE / "assets" / "templates" / "AMET.docx"
LOGO_PATH     = _HERE / "assets" / "logos" / "amet.png"
OUTPUT_DIR    = _HERE / "artifacts"


# ── Domain types ─────────────────────────────────────────────────────────────
class AmetQuestion(BaseModel):
    number: str
    text: str
    mark: int
    btl: str
    co: str


class AmetQuestionPair(BaseModel):
    a: AmetQuestion
    b: AmetQuestion


class AmetPartA(BaseModel):
    total: str
    instruction: str
    questions: list[AmetQuestion]


class AmetPartB(BaseModel):
    total: str
    instruction: str
    question_pairs: list[AmetQuestionPair]


class AmetPartC(BaseModel):
    total: str
    instruction: str
    question: AmetQuestion


class AmetAssessmentRequest(BaseModel):
    exam_type: str
    programme: str
    semester: str
    course_name: str
    course_code: str
    duration: str
    max_marks: int
    instructions: list[str]
    part_a: AmetPartA
    part_b: AmetPartB
    part_c: AmetPartC


# ── Helpers ───────────────────────────────────────────────────────────────────
def _open_template() -> Document:
    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "sectPr":
                body.remove(child)
    else:
        doc = Document()
    return doc


def _insert_logo(doc: Document) -> None:
    logo = LOGO_PATH
    if not logo.exists():
        for ext in (".jpg", ".jpeg", ".PNG", ".JPG"):
            alt = logo.with_suffix(ext)
            if alt.exists():
                logo = alt
                break
        else:
            return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Inches(1.8))


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _bold_center(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_row(table, cells: list[str], bold: bool = False) -> None:
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
        if bold:
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True


# ── Renderer ──────────────────────────────────────────────────────────────────
def _render(request: AmetAssessmentRequest) -> bytes:
    doc = _open_template()
    _insert_logo(doc)

    _bold_center(doc, request.exam_type, size=13)
    doc.add_paragraph(
        f"Programme & Batch: {request.programme}"
        f"{'':>20}Semester : {request.semester}"
    )
    doc.add_paragraph(
        f"Course Name : {request.course_name}"
        f"{'':>20}Course Code : {request.course_code}"
    )
    doc.add_paragraph(
        f"Duration : {request.duration}"
        f"{'':>30}Maximum Marks: {request.max_marks} marks"
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Instructions:").bold = True
    for i, instr in enumerate(request.instructions, start=1):
        doc.add_paragraph(f"{i}. {instr}")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    _set_borders(table)
    for i, h in enumerate(["Question No", "Question", "Mark", "BTL", "CO"]):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = f"PART A ({request.part_a.total}) {request.part_a.instruction}"
    row.cells[0].paragraphs[0].runs[0].bold = True

    for q in request.part_a.questions:
        _add_row(table, [q.number, q.text, str(q.mark), q.btl, q.co])

    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = f"PART B ({request.part_b.total}) {request.part_b.instruction}"
    row.cells[0].paragraphs[0].runs[0].bold = True

    for pair in request.part_b.question_pairs:
        _add_row(table, [pair.a.number, pair.a.text, str(pair.a.mark), pair.a.btl, pair.a.co])
        or_row = table.add_row()
        or_row.cells[0].merge(or_row.cells[4])
        or_row.cells[0].text = "(OR)"
        or_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_row(table, [pair.b.number, pair.b.text, str(pair.b.mark), pair.b.btl, pair.b.co])

    row = table.add_row()
    row.cells[0].merge(row.cells[4])
    row.cells[0].text = f"PART C ({request.part_c.total}) {request.part_c.instruction}"
    row.cells[0].paragraphs[0].runs[0].bold = True

    q = request.part_c.question
    _add_row(table, [q.number, q.text, str(q.mark), q.btl, q.co])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────
def generate(request: AmetAssessmentRequest, filename: str | None = None) -> str:
    """Render the assessment and save it. Returns the output file path."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if filename is None:
        safe = request.course_code.replace(" ", "_")
        filename = f"AMET_{safe}.docx"
    out = OUTPUT_DIR / filename
    out.write_bytes(_render(request))
    return str(out)
