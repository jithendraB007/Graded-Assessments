"""
ANU Graded Assessment Generator
---------------------------------
Self-contained script. All assets are resolved relative to this file.

  assets/templates/ANU.docx  — university page layout template

Usage:
    from generate import generate, AnuAssessmentRequest, ...
    data = AnuAssessmentRequest(...)
    path = generate(data)          # returns saved file path
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pydantic import BaseModel

# ── Paths ─────────────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent
TEMPLATE_PATH = _HERE / "assets" / "templates" / "ANU.docx"
OUTPUT_DIR    = _HERE / "artifacts"


# ── Domain types ──────────────────────────────────────────────────────────────
class AnuSubQuestion(BaseModel):
    sub: str
    text: str
    co: str
    bloom: str


class AnuPartBQuestion(BaseModel):
    number: str
    text: str
    marks: str
    co: str
    bloom: str


class AnuPartA(BaseModel):
    sub_questions: list[AnuSubQuestion]


class AnuPartB(BaseModel):
    questions: list[AnuPartBQuestion]


class AnuAssessmentRequest(BaseModel):
    university_name: str
    batch: str
    exam_type: str
    course_name: str
    date: str
    duration: str
    max_marks: int
    notes: list[str]
    part_a: AnuPartA
    part_b: AnuPartB


# ── Helpers ───────────────────────────────────────────────────────────────────
def _open_template() -> Document:
    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "sectPr":
                body.remove(child)
    else:
        doc = Document()
    return doc


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _bold_center(doc: Document, text: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


# ── Renderer ──────────────────────────────────────────────────────────────────
def _render(request: AnuAssessmentRequest) -> bytes:
    doc = _open_template()

    _bold_center(doc, request.university_name, size=14)
    _bold_center(doc, request.batch)
    _bold_center(doc, request.exam_type)
    _bold_center(doc, request.course_name)
    doc.add_paragraph(
        f"Date: {request.date}   "
        f"Duration: {request.duration}   "
        f"Max. Marks: {request.max_marks}"
    )
    doc.add_paragraph()

    ht_table = doc.add_table(rows=1, cols=12)
    _set_borders(ht_table)
    ht_table.rows[0].cells[0].text = "H.T. No:-"
    doc.add_paragraph()

    for i, note in enumerate(request.notes, start=1):
        doc.add_paragraph(f"{i}. {note}")
    doc.add_paragraph()

    part_a_table = doc.add_table(rows=1, cols=6)
    _set_borders(part_a_table)
    hdr = part_a_table.rows[0]
    for i, h in enumerate(["PART-A", "", "", "", "Course Outcomes", "Bloom's level"]):
        hdr.cells[i].text = h
        if h:
            hdr.cells[i].paragraphs[0].runs[0].bold = True
    hdr.cells[0].merge(hdr.cells[3])

    instr_row = part_a_table.add_row()
    instr_row.cells[0].merge(instr_row.cells[3])
    instr_row.cells[0].text = "Answer all the following short answer questions."

    q_row = part_a_table.add_row()
    q_row.cells[0].text = "1"
    for sq in request.part_a.sub_questions:
        sub_row = part_a_table.add_row()
        sub_row.cells[1].text = sq.sub
        sub_row.cells[2].merge(sub_row.cells[3])
        sub_row.cells[2].text = sq.text
        sub_row.cells[4].text = sq.co
        sub_row.cells[5].text = sq.bloom

    doc.add_paragraph()

    part_b_table = doc.add_table(rows=1, cols=6)
    _set_borders(part_b_table)
    hdr_b = part_b_table.rows[0]
    for i, h in enumerate(["PART B", "", "", "Marks", "Course Outcomes", "Bloom's Level"]):
        hdr_b.cells[i].text = h
        if h:
            hdr_b.cells[i].paragraphs[0].runs[0].bold = True

    for q in request.part_b.questions:
        row = part_b_table.add_row()
        row.cells[0].text = q.number
        row.cells[2].merge(row.cells[1])
        row.cells[1].text = q.text
        row.cells[3].text = q.marks
        row.cells[4].text = q.co
        row.cells[5].text = q.bloom

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────
def generate(request: AnuAssessmentRequest, filename: str | None = None) -> str:
    """Render the assessment and save it. Returns the output file path."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if filename is None:
        safe = request.course_name.replace(" ", "_")[:30]
        filename = f"ANU_{safe}.docx"
    out = OUTPUT_DIR / filename
    out.write_bytes(_render(request))
    return str(out)
