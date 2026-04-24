"""
S-VYASA Graded Assessment Generator
-------------------------------------
Self-contained script. All assets are resolved relative to this file.

  assets/templates/S-Vyasa.docx  — university page layout template
  assets/logos/s-vyasa.jpg       — university logo

Usage:
    from generate import generate, SvyasaAssessmentRequest, ...
    data = SvyasaAssessmentRequest(...)
    path = generate(data)          # returns saved file path
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from pydantic import BaseModel

# ── Paths ─────────────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent
TEMPLATE_PATH = _HERE / "assets" / "templates" / "S-Vyasa.docx"
LOGO_PATH     = _HERE / "assets" / "logos" / "s-vyasa.jpg"
OUTPUT_DIR    = _HERE / "artifacts"


# ── Domain types ──────────────────────────────────────────────────────────────
class SvyasaQuestion(BaseModel):
    number: str
    text: str
    co: str
    rbtl: str
    marks: int


class SvyasaQuestionPair(BaseModel):
    a: SvyasaQuestion
    b: SvyasaQuestion


class SvyasaPartA(BaseModel):
    questions: list[SvyasaQuestion]


class SvyasaPartB(BaseModel):
    question_pairs: list[SvyasaQuestionPair]


class SvyasaAssessmentRequest(BaseModel):
    month_year: str
    academic_year: str
    program: str
    specialization: str
    semester: str
    date_of_exam: str
    course_code: str
    course_name: str
    part_a: SvyasaPartA
    part_b: SvyasaPartB


# ── Helpers ───────────────────────────────────────────────────────────────────
def _open_template() -> Document:
    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "sectPr":
                body.remove(child)
    else:
        doc = Document()
    return doc


def _insert_logo(doc: Document) -> None:
    logo = LOGO_PATH
    if not logo.exists():
        for ext in (".png", ".PNG", ".JPG", ".jpeg"):
            alt = logo.with_suffix(ext)
            if alt.exists():
                logo = alt
                break
        else:
            return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Inches(1.8))


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


# ── Renderer ──────────────────────────────────────────────────────────────────
def _render(request: SvyasaAssessmentRequest) -> bytes:
    doc = _open_template()
    _insert_logo(doc)

    usn_table = doc.add_table(rows=1, cols=12)
    _set_borders(usn_table)
    usn_table.rows[0].cells[0].text = "USN"
    doc.add_paragraph()

    hdr_table = doc.add_table(rows=5, cols=4)
    _set_borders(hdr_table)
    hdr_data = [
        ["Month & Year of Examination", request.month_year, "Academic year", request.academic_year],
        ["Program", request.program, "Specialization", request.specialization],
        ["Semester", request.semester, "Date of Examination", request.date_of_exam],
        ["Course Code", request.course_code, request.course_code, request.course_code],
        ["Course Name", request.course_name, request.course_name, request.course_name],
    ]
    for r_idx, row_data in enumerate(hdr_data):
        for c_idx, text in enumerate(row_data):
            hdr_table.rows[r_idx].cells[c_idx].text = text
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Part – A").bold = True

    part_a_count = len(request.part_a.questions)
    marks_each = request.part_a.questions[0].marks if request.part_a.questions else 3
    doc.add_paragraph(
        f"Answer all the questions   "
        f"{part_a_count} Q x {marks_each} M = {part_a_count * marks_each}"
    )

    table_a = doc.add_table(rows=1, cols=5)
    _set_borders(table_a)
    for i, h in enumerate(["Q.No.", "Questions", "CO", "RBTL", "Marks"]):
        table_a.rows[0].cells[i].text = h
        table_a.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for q in request.part_a.questions:
        row = table_a.add_row()
        row.cells[0].text = q.number
        row.cells[1].text = q.text
        row.cells[2].text = q.co
        row.cells[3].text = q.rbtl
        row.cells[4].text = str(q.marks)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Part - B").bold = True

    pair_count = len(request.part_b.question_pairs)
    marks_each_b = request.part_b.question_pairs[0].a.marks if request.part_b.question_pairs else 14
    doc.add_paragraph(
        f"Answer all the questions   "
        f"{pair_count} Q x {marks_each_b} M = {pair_count * marks_each_b}"
    )

    table_b = doc.add_table(rows=1, cols=5)
    _set_borders(table_b)
    for i, h in enumerate(["Q.No.", "Questions", "CO", "RBTL", "Marks"]):
        table_b.rows[0].cells[i].text = h
        table_b.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for pair in request.part_b.question_pairs:
        row_a = table_b.add_row()
        row_a.cells[0].text = pair.a.number
        row_a.cells[1].text = pair.a.text
        row_a.cells[2].text = pair.a.co
        row_a.cells[3].text = pair.a.rbtl
        row_a.cells[4].text = str(pair.a.marks)

        or_row = table_b.add_row()
        or_row.cells[0].merge(or_row.cells[4])
        or_row.cells[0].text = "OR"
        or_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_b = table_b.add_row()
        row_b.cells[0].text = pair.b.number
        row_b.cells[1].text = pair.b.text
        row_b.cells[2].text = pair.b.co
        row_b.cells[3].text = pair.b.rbtl
        row_b.cells[4].text = str(pair.b.marks)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────
def generate(request: SvyasaAssessmentRequest, filename: str | None = None) -> str:
    """Render the assessment and save it. Returns the output file path."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if filename is None:
        safe = request.course_code.replace(" ", "_")
        filename = f"SVYASA_{safe}.docx"
    out = OUTPUT_DIR / filename
    out.write_bytes(_render(request))
    return str(out)
