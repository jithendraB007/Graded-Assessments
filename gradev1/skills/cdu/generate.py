"""
CDU Graded Assessment Generator
---------------------------------
Self-contained script. All assets are resolved relative to this file.

  assets/templates/CDU.docx  — university page layout template

Usage:
    from generate import generate, CduAssessmentRequest, ...
    data = CduAssessmentRequest(...)
    path = generate(data)          # returns saved file path
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel

# ── Paths ─────────────────────────────────────────────────────────────────────
_HERE = Path(__file__).parent
TEMPLATE_PATH = _HERE / "assets" / "templates" / "CDU.docx"
OUTPUT_DIR    = _HERE / "artifacts"


# ── Domain types ──────────────────────────────────────────────────────────────
class CduQuestion(BaseModel):
    number: str
    text: str


class CduQuestionPair(BaseModel):
    a: CduQuestion
    b: CduQuestion


class CduSectionA(BaseModel):
    instruction: str
    questions: list[CduQuestion]


class CduSectionB(BaseModel):
    instruction: str
    question_pairs: list[CduQuestionPair]


class CduSet(BaseModel):
    label: str
    section_a: CduSectionA
    section_b: CduSectionB


class CduAssessmentRequest(BaseModel):
    university_name: str
    course_info: str
    time: str
    max_marks: int
    sets: list[CduSet]


# ── Helpers ───────────────────────────────────────────────────────────────────
def _open_template() -> Document:
    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "sectPr":
                body.remove(child)
    else:
        doc = Document()
    return doc


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _render_set(doc: Document, cdu_set: CduSet, university_name: str, course_info: str, time: str, max_marks: int) -> None:
    table = doc.add_table(rows=0, cols=2)
    _set_borders(table)

    hdr_row = table.add_row()
    hdr_row.cells[0].merge(hdr_row.cells[1])
    hdr_row.cells[0].text = f"{cdu_set.label}\n{university_name}\n{course_info}"
    hdr_row.cells[0].paragraphs[0].runs[0].bold = True
    hdr_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    time_row = table.add_row()
    time_row.cells[0].text = f"Time: {time}]"
    time_row.cells[1].text = f"[Max. Marks: {max_marks}"

    sec_a_row = table.add_row()
    sec_a_row.cells[0].merge(sec_a_row.cells[1])
    sec_a_row.cells[0].text = "Section - A"
    sec_a_row.cells[0].paragraphs[0].runs[0].bold = True

    instr_row = table.add_row()
    instr_row.cells[0].merge(instr_row.cells[1])
    instr_row.cells[0].text = cdu_set.section_a.instruction

    for q in cdu_set.section_a.questions:
        row = table.add_row()
        row.cells[0].text = q.number
        row.cells[1].text = q.text

    sec_b_row = table.add_row()
    sec_b_row.cells[0].merge(sec_b_row.cells[1])
    sec_b_row.cells[0].text = "Section - B"
    sec_b_row.cells[0].paragraphs[0].runs[0].bold = True

    instr_b_row = table.add_row()
    instr_b_row.cells[0].merge(instr_b_row.cells[1])
    instr_b_row.cells[0].text = cdu_set.section_b.instruction

    for pair in cdu_set.section_b.question_pairs:
        row_a = table.add_row()
        row_a.cells[0].text = pair.a.number
        row_a.cells[1].text = pair.a.text

        or_row = table.add_row()
        or_row.cells[0].merge(or_row.cells[1])
        or_row.cells[0].text = "OR"
        or_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_b = table.add_row()
        row_b.cells[0].text = pair.b.number
        row_b.cells[1].text = pair.b.text


# ── Renderer ──────────────────────────────────────────────────────────────────
def _render(request: CduAssessmentRequest) -> bytes:
    doc = _open_template()

    for i, cdu_set in enumerate(request.sets):
        _render_set(doc, cdu_set, request.university_name, request.course_info, request.time, request.max_marks)
        if i < len(request.sets) - 1:
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────
def generate(request: CduAssessmentRequest, filename: str | None = None) -> str:
    """Render the assessment and save it. Returns the output file path."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if filename is None:
        safe = request.course_info.replace(" ", "_")[:30]
        filename = f"CDU_{safe}.docx"
    out = OUTPUT_DIR / filename
    out.write_bytes(_render(request))
    return str(out)
