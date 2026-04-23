"""
One-time script to generate assets/templates/reference-assessment-template.docx.

Users copy this file, open it in Word, add their university logo and branding,
then save it as assets/templates/{university_id}-assessment.docx.

The docxtpl Jinja2 placeholders embedded as plain text are left intact by Word
as long as users don't reformat or split the placeholder paragraphs.

Run from the project root:
    python scripts/create_assessment_reference_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def _center(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("=" * 60)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def create_reference_template(output_path: Path) -> None:
    doc = Document()

    # ─── COVER PAGE ──────────────────────────────────────────────
    # Placeholder for university logo — user inserts image here in Word
    p = doc.add_paragraph()
    _center(p)
    run = p.add_run("[INSERT UNIVERSITY LOGO HERE]")
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.italic = True

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    _center(p)
    run = p.add_run("{{ cover.university_name }}")
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    _center(p)
    p.add_run("{{ cover.department }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    _center(p)
    p.add_run("{{ cover.course_name }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    _center(p)
    run = p.add_run("{{ cover.assessment_title }}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    _center(p)
    p.add_run("Semester: {{ cover.semester }}")

    p = doc.add_paragraph()
    _center(p)
    p.add_run("Date: {{ cover.date }}")

    doc.add_page_break()

    # ─── INSTRUCTIONS ────────────────────────────────────────────
    doc.add_heading("INSTRUCTIONS", level=1)

    doc.add_paragraph("Duration:        {{ instructions.duration }}")
    doc.add_paragraph("Total Marks:     {{ instructions.total_marks }}")
    doc.add_paragraph("Passing Marks:   {{ instructions.passing_marks }}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Attempt Rules:")
    run.bold = True

    # Paragraph-level Jinja2 loop — docxtpl expands these paragraphs per item
    doc.add_paragraph("{%p for rule in instructions.attempt_rules %}")
    doc.add_paragraph("  • {{ rule }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("General Rules:")
    run.bold = True

    doc.add_paragraph("{%p for note in instructions.general_notes %}")
    doc.add_paragraph("  • {{ note }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_page_break()

    # ─── QUESTIONS ───────────────────────────────────────────────
    # Outer loop: sections
    doc.add_paragraph("{%p for section in sections %}")

    _add_separator(doc)
    p = doc.add_paragraph()
    run = p.add_run(
        "SECTION {{ section.letter }}: {{ section.type_name }}"
        "  [{{ section.total_marks }} Marks]"
    )
    run.bold = True
    _add_separator(doc)

    doc.add_paragraph()

    # Inner loop: questions within each section
    doc.add_paragraph("{%p for q in section.questions %}")

    doc.add_paragraph("Q{{ q.number }}. {{ q.text }}  [{{ q.marks }} marks]")

    # MCQ options — only rendered when question type is mcq
    doc.add_paragraph("{%p if q.type == 'mcq' %}")
    doc.add_paragraph("   A. {{ q.options.a }}")
    doc.add_paragraph("   B. {{ q.options.b }}")
    doc.add_paragraph("   C. {{ q.options.c }}")
    doc.add_paragraph("   D. {{ q.options.d }}")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()  # blank line between questions

    doc.add_paragraph("{%p endfor %}")  # end questions loop

    doc.add_paragraph("{%p endfor %}")  # end sections loop

    # ─── SAVE ────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Reference template created: {output_path}")
    print()
    print("Next steps:")
    print("  1. Open the file in Microsoft Word or LibreOffice")
    print("  2. Replace [INSERT UNIVERSITY LOGO HERE] with your logo image")
    print("  3. Apply university branding (fonts, colors, header/footer)")
    print("  4. Save a copy as: assets/templates/{university_id}-assessment.docx")
    print("  5. Do NOT edit the {{ }} or {%p %} placeholder text")


if __name__ == "__main__":
    project_root = Path(__file__).parent.parent
    output = project_root / "assets" / "templates" / "reference-assessment-template.docx"
    create_reference_template(output)
